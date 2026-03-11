from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins
section = doc.sections[0]
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Color palette
BLACK   = RGBColor(0x11, 0x11, 0x11)
GREEN   = RGBColor(0x16, 0xA3, 0x4A)
GRAY    = RGBColor(0x6B, 0x72, 0x80)
LGRAY   = RGBColor(0xD1, 0xD5, 0xDB)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
CREAM   = RGBColor(0xF7, 0xF6, 0xF2)

# ── Helper: set run font
def run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

# ── Helper: paragraph spacing
def para_space(p, before=0, after=6):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)

# ── Helper: shading a paragraph
def shade_para(para, fill_hex="F7F6F2"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

# ── Helper: horizontal rule
def add_rule(doc):
    p = doc.add_paragraph()
    para_space(p, 4, 4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E5E7EB')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════
# COVER BLOCK
# ══════════════════════════════════════
p = doc.add_paragraph()
para_space(p, 0, 2)
r = p.add_run("CMO MARKETING SUITE")
run_font(r, "Calibri", 9, bold=True, color=GREEN)
r.font.all_caps = True

p = doc.add_paragraph()
para_space(p, 0, 8)
r = p.add_run("Creating a Legit Peptide Company Is the Easy Part.\nCompeting with Silicon Valley Giants Is the Hard Part.")
run_font(r, "Calibri", 28, bold=True, color=BLACK)
p.paragraph_format.line_spacing = Pt(32)

p = doc.add_paragraph()
para_space(p, 0, 4)
r = p.add_run("Until Now.")
run_font(r, "Calibri", 28, bold=True, italic=True, color=GREEN)

p = doc.add_paragraph()
para_space(p, 8, 24)
r = p.add_run("Plug & Play Peptides  ·  plugandplaypeptides.com  ·  help@plugandplaypeptides.com  ·  415-619-7661")
run_font(r, "Calibri", 9, color=GRAY)

add_rule(doc)

# ══════════════════════════════════════
# SECTION 1: THE PROBLEM
# ══════════════════════════════════════
p = doc.add_paragraph()
para_space(p, 16, 4)
r = p.add_run("THE PROBLEM")
run_font(r, "Calibri", 8, bold=True, color=GREEN)
r.font.all_caps = True

p = doc.add_paragraph()
para_space(p, 0, 6)
r = p.add_run("They don't just have more money. They have a better machine.")
run_font(r, "Calibri", 22, bold=True, color=BLACK)
p.paragraph_format.line_spacing = Pt(26)

p = doc.add_paragraph()
para_space(p, 0, 16)
r = p.add_run(
    "The brands you're competing against — Hims & Hers, Roman, Ro Health, Noom Med, Found, Calibrate — "
    "didn't win because of better products. They won because of infrastructure you don't have yet."
)
run_font(r, "Calibri", 11, color=GRAY)

problems = [
    (
        "01  They know you're leaving before you do.",
        "Predictive models built on millions of patient events fire re-engagement sequences weeks before a patient "
        "considers canceling. That window — your second chance — fires automatically. Yours doesn't."
    ),
    (
        "02  Their brand lives in patients' heads rent-free.",
        "Every patient who orders from Hims enters a post-purchase experience so well-designed that the brand "
        "becomes part of their identity. You ship a box. They build a relationship."
    ),
    (
        "03  Their revenue compounds. Yours resets every month.",
        "When every retained patient fuels acquisition, every protocol completion unlocks an upsell, and every "
        "referral compounds the base — revenue doesn't grow linearly. It accelerates. That's what you're up against."
    ),
]

for heading, body in problems:
    p = doc.add_paragraph()
    para_space(p, 10, 2)
    r = p.add_run(heading)
    run_font(r, "Calibri", 11, bold=True, color=BLACK)

    p = doc.add_paragraph()
    para_space(p, 0, 12)
    r = p.add_run(body)
    run_font(r, "Calibri", 10.5, color=GRAY)

add_rule(doc)

# ══════════════════════════════════════
# SECTION 2: THE CMO SUITE REVEAL
# ══════════════════════════════════════
p = doc.add_paragraph()
para_space(p, 16, 4)
r = p.add_run("THE CMO MARKETING SUITE")
run_font(r, "Calibri", 8, bold=True, color=GREEN)
r.font.all_caps = True

p = doc.add_paragraph()
para_space(p, 0, 6)
r = p.add_run("What took Hims $100M and five years to build — yours on day one.")
run_font(r, "Calibri", 22, bold=True, color=BLACK)
p.paragraph_format.line_spacing = Pt(26)

p = doc.add_paragraph()
para_space(p, 0, 16)
r = p.add_run(
    "Not software. Not a dashboard. A complete growth machine — retention, acquisition, conversion, and "
    "intelligence — running under your brand, without a single hire."
)
run_font(r, "Calibri", 11, color=GRAY)

# 4-column benefit table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells

benefit_cols = [
    ("RETENTION", "Patients who never leave.", 
     "Check-ins, protocol updates, reorder signals — all timed precisely. Your patients feel cared for because they are."),
    ("ACQUISITION", "New patients who cost less to find.", 
     "Every retained patient becomes a referral engine. Acquisition compounds the longer you run — without spending more."),
    ("REVENUE", "Money that arrives while you sleep.", 
     "Protocol progressions. Smart reorders. Tiered plans. Every patient journey is designed to extend naturally."),
    ("INTELLIGENCE", "You know what's working before your competitors notice.", 
     "Cohort data, LTV signals, churn patterns — surfaced and acted on automatically. No data team required."),
]

for i, (tag, desire, body) in enumerate(benefit_cols):
    cell = hdr_cells[i]
    cell.width = Inches(1.5)
    # Tag
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(tag)
    run_font(r, "Calibri", 7.5, bold=True, color=GREEN)
    r.font.all_caps = True
    para_space(p, 6, 4)
    # Desire
    p2 = cell.add_paragraph()
    r2 = p2.add_run(desire)
    run_font(r2, "Calibri", 10.5, bold=True, color=BLACK)
    para_space(p2, 0, 4)
    # Body
    p3 = cell.add_paragraph()
    r3 = p3.add_run(body)
    run_font(r3, "Calibri", 9, color=GRAY)
    para_space(p3, 0, 6)

# Remove table borders except internal lines
from docx.oxml import OxmlElement
tbl = table._tbl
tblPr = tbl.find(qn('w:tblPr'))
if tblPr is None:
    tblPr = OxmlElement('w:tblPr')
    tbl.insert(0, tblPr)
tblBorders = OxmlElement('w:tblBorders')
for border_name in ['top','left','bottom','right','insideH','insideV']:
    b = OxmlElement(f'w:{border_name}')
    if border_name in ['insideH', 'insideV']:
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '2')
        b.set(qn('w:color'), 'E5E7EB')
    else:
        b.set(qn('w:val'), 'none')
    tblBorders.append(b)
tblPr.append(tblBorders)

p = doc.add_paragraph()
para_space(p, 6, 0)

add_rule(doc)

# ══════════════════════════════════════
# SECTION 3: WHAT IT FEELS LIKE
# ══════════════════════════════════════
p = doc.add_paragraph()
para_space(p, 16, 4)
r = p.add_run("WHAT IT FEELS LIKE")
run_font(r, "Calibri", 8, bold=True, color=GREEN)
r.font.all_caps = True

p = doc.add_paragraph()
para_space(p, 0, 6)
r = p.add_run('You wake up. Your brand grew while you slept.')
run_font(r, "Calibri", 20, bold=True, italic=True, color=BLACK)
p.paragraph_format.line_spacing = Pt(24)

p = doc.add_paragraph()
para_space(p, 0, 14)
r = p.add_run("Every morning your brand has been working overnight. Not metaphorically. Literally.")
run_font(r, "Calibri", 11, color=GRAY)

feed_events = [
    ("RETENTION",    "A patient who ordered 18 days ago received a reorder prompt at exactly the right moment. They're checking out now."),
    ("LTV",          "A patient completed their first semaglutide protocol. A physician-guided progression to the next tier started automatically."),
    ("WIN-BACK",     "A patient who went quiet at day 12 received a personalized check-in. They replied, and reordered."),
    ("REFERRAL",     "A patient referred two friends. Both completed intake. All three patients now carry your brand."),
    ("INTELLIGENCE", "Cohort data flagged that tirzepatide patients have 22% higher 90-day LTV in your brand. Logged for review."),
]

for label, text in feed_events:
    p = doc.add_paragraph()
    para_space(p, 4, 4)
    r1 = p.add_run(f"[{label}]  ")
    run_font(r1, "Calibri", 8, bold=True, color=GREEN)
    r1.font.all_caps = True
    r2 = p.add_run(text)
    run_font(r2, "Calibri", 10.5, color=BLACK)
    shade_para(p, "F7F6F2")

p = doc.add_paragraph()
para_space(p, 6, 0)
r = p.add_run("All of this. While you were doing something else.")
run_font(r, "Calibri", 9, italic=True, color=GRAY)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_rule(doc)

# ══════════════════════════════════════
# SECTION 4: PARTNER RESULTS / PROOF
# ══════════════════════════════════════
p = doc.add_paragraph()
para_space(p, 16, 4)
r = p.add_run("PARTNER RESULTS")
run_font(r, "Calibri", 8, bold=True, color=GREEN)
r.font.all_caps = True

p = doc.add_paragraph()
para_space(p, 0, 16)
r = p.add_run("What changes when you stop competing alone.")
run_font(r, "Calibri", 22, bold=True, color=BLACK)
p.paragraph_format.line_spacing = Pt(26)

testimonials = [
    (
        "★★★★★",
        "Month one I made $22,400. But what actually surprised me was month three — I hadn't run a single new ad. "
        "The revenue came from patients who already ordered.",
        "Alex M.",
        "Fitness Creator · 800K followers",
        "$22,400 month one"
    ),
    (
        "★★★★★",
        "I own a med spa. I know what a real patient experience looks like. What my brand delivers through "
        "Plug & Play is better than most clinics I've seen — and I'm not managing any of it.",
        "Dr. Sarah K.",
        "Med Spa Owner · Miami, FL",
        "$45K/mo recurring"
    ),
    (
        "★★★★★",
        "I've promoted brands for years. This is the first time I've promoted something where the brand itself "
        "keeps the relationship alive after I send someone. My audience trusts me more, not less.",
        "Jordan T.",
        "Biohacking Creator · 1.2M followers",
        "Audience trust preserved"
    ),
]

for stars, quote, name, role, tag in testimonials:
    p = doc.add_paragraph()
    para_space(p, 10, 2)
    r = p.add_run(stars + "  ")
    run_font(r, "Calibri", 11, color=RGBColor(0xF5, 0x9E, 0x0B))
    r2 = p.add_run(tag)
    run_font(r2, "Calibri", 8, bold=True, color=GREEN)

    p = doc.add_paragraph()
    para_space(p, 2, 2)
    r = p.add_run(f'"{quote}"')
    run_font(r, "Calibri", 11, italic=True, color=BLACK)

    p = doc.add_paragraph()
    para_space(p, 0, 12)
    r = p.add_run(f"{name}  —  {role}")
    run_font(r, "Calibri", 9, bold=True, color=GRAY)

add_rule(doc)

# ══════════════════════════════════════
# SECTION 5: CTA / WHAT YOU WAKE UP WITH
# ══════════════════════════════════════
p = doc.add_paragraph()
para_space(p, 16, 4)
r = p.add_run("ONE PARTNER PER NICHE")
run_font(r, "Calibri", 8, bold=True, color=GREEN)
r.font.all_caps = True

p = doc.add_paragraph()
para_space(p, 0, 6)
r = p.add_run("The hard part just got a lot easier.")
run_font(r, "Calibri", 26, bold=True, color=BLACK)
p.paragraph_format.line_spacing = Pt(30)

p = doc.add_paragraph()
para_space(p, 0, 16)
r = p.add_run(
    "The CMO Suite is included with every Plug & Play partnership. The compliance infrastructure, the physician "
    "network, fulfillment, and the full growth machine — all under your brand. We accept one partner per niche "
    "and geography. The brands applying today are the ones who will own this space."
)
run_font(r, "Calibri", 11, color=GRAY)

p = doc.add_paragraph()
para_space(p, 4, 8)
r = p.add_run("What you wake up with:")
run_font(r, "Calibri", 12, bold=True, color=BLACK)

outcomes = [
    ("A physician-supervised brand that passes any audit.",
     "FDA-registered pharmacies · Licensed MDs in all 50 states · HIPAA-compliant platform"),
    ("Patients who feel like they're cared for by a real brand.",
     "Onboarding · Check-ins · Reorders · Protocol progressions — all automatic, all under your brand"),
    ("Revenue that grows without proportional effort.",
     "Retention engine · Referral loops · Protocol upsells — compounding month over month"),
    ("Intelligence that gets smarter every month.",
     "LTV cohorts · Churn signals · Acquisition data — surfaced and acted on automatically"),
    ("Live in 2–4 weeks. No staff. No overhead.",
     "We build everything. You promote your brand. We handle the rest."),
]

for desire, sub in outcomes:
    p = doc.add_paragraph(style='List Bullet')
    para_space(p, 4, 2)
    r1 = p.add_run(desire + "  ")
    run_font(r1, "Calibri", 11, bold=True, color=BLACK)
    r2 = p.add_run(sub)
    run_font(r2, "Calibri", 10, color=GRAY)

add_rule(doc)

# ══════════════════════════════════════
# FOOTER / CONTACT BLOCK
# ══════════════════════════════════════
p = doc.add_paragraph()
para_space(p, 16, 4)
r = p.add_run("APPLY TO PARTNER  —  PLUGANDPLAYPEPTIDES.COM/APPLY")
run_font(r, "Calibri", 8, bold=True, color=GREEN)
r.font.all_caps = True

p = doc.add_paragraph()
para_space(p, 0, 4)
r = p.add_run("5-minute application · 100% confidential · No commitment required")
run_font(r, "Calibri", 11, italic=True, color=GRAY)

p = doc.add_paragraph()
para_space(p, 8, 4)
r = p.add_run("Contact Us")
run_font(r, "Calibri", 10, bold=True, color=BLACK)

contact_lines = [
    "help@plugandplaypeptides.com",
    "415-619-7661",
    "Mon–Fri, 9am–5pm MT",
    "1001 S Main St, Suite 12636, Kalispell, MT 59901",
    "© 2026 Plug and Play Peptides, LLC  ·  A Montana Limited Liability Company",
]
for line in contact_lines:
    p = doc.add_paragraph()
    para_space(p, 0, 2)
    r = p.add_run(line)
    run_font(r, "Calibri", 9, color=GRAY)

# ── Save
out_path = "/home/user/webapp/CMO_Marketing_Suite_Plug_and_Play_Peptides.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
